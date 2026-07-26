import os
import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    in_code_block = False
    code_lines = []

    in_table = False
    table_lines = []

    def flush_table(t_lines):
        if not t_lines:
            return
        rows_data = []
        for line in t_lines:
            line_str = line.strip()
            if line_str.startswith('|') and line_str.endswith('|'):
                parts = [p.strip() for p in line_str.split('|')[1:-1]]
                # Skip divider row like |---|---|
                if all(re.match(r'^:?-+:?$', p) for p in parts if p):
                    continue
                rows_data.append(parts)
        
        if not rows_data:
            return
        
        col_count = max(len(r) for r in rows_data)
        table = doc.add_table(rows=len(rows_data), cols=col_count)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        for r_idx, row in enumerate(rows_data):
            for c_idx, val in enumerate(row):
                if c_idx < col_count:
                    cell = table.cell(r_idx, c_idx)
                    cell.text = val
                    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
                    if r_idx == 0:
                        set_cell_background(cell, "2D3748") # Dark gray header
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.bold = True
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.name = 'Malgun Gothic'
                    else:
                        bg_color = "F7FAFC" if r_idx % 2 == 1 else "FFFFFF"
                        set_cell_background(cell, bg_color)
                        for p in cell.paragraphs:
                            for run in p.runs:
                                run.font.color.rgb = RGBColor(45, 55, 72)
                                run.font.name = 'Malgun Gothic'
                                run.font.size = Pt(9.5)
        doc.add_paragraph() # spacing

    for line in lines:
        stripped = line.strip()

        # Handle Code Block ```
        if stripped.startswith('```'):
            if in_code_block:
                # Flush code block
                code_text = "".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(4)
                p.paragraph_format.space_after = Pt(4)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(45, 55, 72)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            continue

        if in_code_block:
            code_lines.append(line)
            continue

        # Handle Tables
        if stripped.startswith('|') and stripped.endswith('|'):
            in_table = True
            table_lines.append(line)
            continue
        else:
            if in_table:
                flush_table(table_lines)
                table_lines = []
                in_table = False

        if not stripped:
            continue

        # Headers
        if stripped.startswith('# '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(stripped[2:])
            run.font.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(26, 54, 93) # Deep Blue
            run.font.name = 'Malgun Gothic'
        elif stripped.startswith('## '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[3:])
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(43, 108, 176)
            run.font.name = 'Malgun Gothic'
        elif stripped.startswith('### '):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(stripped[4:])
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(45, 55, 72)
            run.font.name = 'Malgun Gothic'
        elif stripped.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(stripped[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(113, 128, 150)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(10)
        elif stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped):
            p = doc.add_paragraph(style='List Bullet' if not re.match(r'^\d+\.\s', stripped) else 'List Number')
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            text_content = re.sub(r'^[-\*\d\.]+\s*', '', stripped)
            
            # Simple bold parsing **text**
            parts = re.split(r'(\*\*.*?\*\*)', text_content)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Malgun Gothic'
                r.font.size = Pt(10)
        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(4)
            parts = re.split(r'(\*\*.*?\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.font.bold = True
                else:
                    r = p.add_run(part)
                r.font.name = 'Malgun Gothic'
                r.font.size = Pt(10)

    if in_table:
        flush_table(table_lines)

    doc.save(docx_path)
    print(f"Successfully converted: {md_path} -> {docx_path}")

def convert_all_in_dir(root_dir):
    converted_files = []
    for dirpath, _, filenames in os.walk(root_dir):
        for fname in filenames:
            if fname.endswith('.md'):
                md_path = os.path.join(dirpath, fname)
                docx_path = os.path.splitext(md_path)[0] + '.docx'
                try:
                    md_to_docx(md_path, docx_path)
                    converted_files.append(docx_path)
                except Exception as e:
                    print(f"Error converting {md_path}: {e}")
    return converted_files

if __name__ == "__main__":
    vault_dir = r"c:\Users\master\agrolib"
    files = convert_all_in_dir(vault_dir)
    print(f"\nTotal {len(files)} docx files generated successfully.")
